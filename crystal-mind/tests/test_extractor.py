"""
Offline tests for the content extractor — covers text, docx, pdf, unknown
formats, truncation, and error paths without any network or API keys.
"""

from __future__ import annotations

from pathlib import Path

import pytest

from crystal_mind.collector.extractor import extract


@pytest.mark.offline
def test_extract_text_formats(tmp_path: Path):
    for ext in (".md", ".txt", ".py", ".json"):
        p = tmp_path / f"f{ext}"
        p.write_text("hello world", encoding="utf-8")
        assert extract(p) == "hello world"


@pytest.mark.offline
def test_extract_unknown_extension_returns_empty(tmp_path: Path):
    p = tmp_path / "image.png"
    p.write_bytes(b"\x89PNG\r\n")
    assert extract(p) == ""


@pytest.mark.offline
def test_extract_missing_file_returns_empty(tmp_path: Path):
    # OSError path in _read_text
    assert extract(tmp_path / "does_not_exist.txt") == ""


@pytest.mark.offline
def test_extract_truncates_to_max_chars(tmp_path: Path):
    p = tmp_path / "big.txt"
    p.write_text("x" * 10_000, encoding="utf-8")
    assert len(extract(p, max_chars=100)) == 100


@pytest.mark.offline
def test_extract_docx(tmp_path: Path):
    from docx import Document

    doc = Document()
    doc.add_paragraph("first paragraph")
    doc.add_paragraph("second paragraph")
    p = tmp_path / "doc.docx"
    doc.save(str(p))

    out = extract(p)
    assert "first paragraph" in out
    assert "second paragraph" in out


@pytest.mark.offline
def test_extract_pdf(tmp_path: Path):
    import fitz  # pymupdf

    doc = fitz.open()
    page = doc.new_page()
    page.insert_text((72, 72), "pdf sample text")
    p = tmp_path / "doc.pdf"
    doc.save(str(p))
    doc.close()

    assert "pdf sample text" in extract(p)


@pytest.mark.offline
def test_extract_corrupt_pdf_returns_empty(tmp_path: Path):
    # Exception path in _read_pdf (not a valid PDF)
    p = tmp_path / "broken.pdf"
    p.write_bytes(b"not a real pdf")
    assert extract(p) == ""
